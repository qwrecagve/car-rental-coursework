from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Sarlavha
    title = doc.add_heading('AVTOMOBIL IJARASI TIZIMI - LOYIHA HISOBOTI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Mundarija
    doc.add_heading('Mundarija', level=1)
    doc.add_paragraph('1. Kirish\n2. Ishlatilgan texnologiyalar\n3. Loyiha strukturasi\n4. Backend tahlili\n5. Frontend tahlili\n6. Azure SQL integratsiyasi\n7. Deployment bosqichlari\n8. Xulosa')

    # 1. Kirish
    doc.add_heading('1. Kirish', level=1)
    doc.add_paragraph('Ushbu loyiha zamonaviy avtomobil ijarasi xizmatini raqamlashtirish maqsadida yaratilgan. Tizim orqali mijozlar bo\'sh mashinalarni ko\'rishlari, ijaraga olishlari va qaytarib berishlari mumkin. Barcha ma\'lumotlar xavfsiz ravishda bulutli platformada saqlanadi.')

    # 2. Texnologiyalar
    doc.add_heading('2. Ishlatilgan texnologiyalar', level=1)
    doc.add_paragraph('Loyiha quyidagi texnologiyalar asosida qurilgan:', style='List Bullet')
    doc.add_paragraph('Backend: Python 3.10 + FastAPI', style='List Bullet')
    doc.add_paragraph('Frontend: HTML5, Vanilla CSS, JavaScript', style='List Bullet')
    doc.add_paragraph('Database: Microsoft Azure SQL', style='List Bullet')
    doc.add_paragraph('Drayver: pyodbc (Azure SQL ulanishi uchun)', style='List Bullet')

    # 3. Backend tahlili
    doc.add_heading('3. Backend tahlili', level=1)
    doc.add_paragraph('Backend qismi FastAPI ramkasida yozilgan bo\'lib, u juda yuqori tezlikda ishlaydi. Ma\'lumotlar bazasi bilan ishlash uchun system.py fayli mas\'uldir.')
    
    # Kod namunasi
    doc.add_heading('Backend: system.py (Ulanish qismi)', level=2)
    code = doc.add_paragraph()
    run = code.add_run('def get_db_connection():\n    conn_str = os.getenv("AZURE_SQL_CONNECTIONSTRING")\n    conn_str = conn_str.replace("Encrypt=True", "Encrypt=yes")\n    return pyodbc.connect(conn_str)')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # 4. Azure SQL integratsiyasi
    doc.add_heading('4. Azure SQL integratsiyasi', level=1)
    doc.add_paragraph('Azure SQL bulutli bazasidan foydalanish loyihaning xavfsizligini ta\'minlaydi. Quyidagi amallar bajarildi:')
    doc.add_paragraph('Firewall qoidalari o\'rnatildi (App Service ulanishi uchun).', style='List Bullet')
    doc.add_paragraph('Connection String Azure Portalda "Configuration" bo\'limiga qo\'shildi.', style='List Bullet')
    doc.add_paragraph('Jadvallar sxemasi pyodbc yordamida dinamik yaratildi.', style='List Bullet')

    # 5. Deployment
    doc.add_heading('5. Deployment bosqichlari', level=1)
    doc.add_paragraph('Loyiha Azure App Service platformasiga GitHub Actions orqali joylashtirildi. Bu har bir koddagi o\'zgarishni avtomatik ravishda serverga yetkazish imkonini beradi.')

    # Xulosa
    doc.add_heading('6. Xulosa', level=1)
    doc.add_paragraph('Loyiha yakunida to\'liq ishlaydigan, bulutli texnologiyalar bilan integratsiya qilingan zamonaviy tizim olindi. Tizim istalgan vaqtda kengaytirilishi va yangi funksiyalar qo\'shilishi mumkin.')

    # Faylni saqlash
    doc.save('Loyiha_Hujjati.docx')
    print("Fayl muvaffaqiyatli yaratildi: Loyiha_Hujjati.docx")

if __name__ == "__main__":
    create_report()
